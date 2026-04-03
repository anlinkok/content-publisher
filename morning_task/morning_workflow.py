"""
晨间新闻自动化工作流
流程：启动代理 -> 搜索新闻 -> 改写文章 -> 保存桌面 -> 关闭代理 -> 定时发布
"""

import asyncio
import subprocess
import os
import sys
import json
import time
from datetime import datetime, timedelta
from pathlib import Path

# 配置
PROXY_LNK = r"C:\Users\Public\Desktop\精灵学院.lnk"
DESKTOP_PATH = r"C:\Users\Administrator\Desktop"
CONTENT_PUBLISHER_PATH = r"D:\kimijiaoben\content-publisher"
TOPICS = ["全球变暖", "人工智能", "新能源", "科技动态", "国际新闻"]


def log(msg):
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_msg = f"[{timestamp}] {msg}"
    print(log_msg)
    
    # 写入日志文件
    log_file = Path(DESKTOP_PATH) / "morning_task.log"
    try:
        with open(log_file, "a", encoding="utf-8") as f:
            f.write(f"{log_msg}\n")
    except:
        pass


def start_proxy():
    """启动代理软件"""
    log("启动精灵学院代理...")
    try:
        # 使用start命令启动快捷方式
        subprocess.Popen(
            f'start "" "{PROXY_LNK}"',
            shell=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )
        log("代理启动命令已发送")
        time.sleep(10)  # 等待代理连接
        return True
    except Exception as e:
        log(f"启动代理失败: {e}")
        return False


def close_proxy():
    """关闭代理进程"""
    log("关闭精灵学院代理...")
    try:
        # 尝试关闭常见代理进程名
        process_names = ["精灵学院", "v2ray", "clash", "shadowsocks", "trojan"]
        for name in process_names:
            try:
                subprocess.run(
                    f"taskkill /f /im {name}.exe",
                    shell=True,
                    capture_output=True,
                    timeout=5
                )
            except:
                pass
        log("代理进程已关闭")
        return True
    except Exception as e:
        log(f"关闭代理失败: {e}")
        return False


def search_and_rewrite():
    """搜索新闻并改写文章"""
    import random
    
    topic = random.choice(TOPICS)
    log(f"选择主题: {topic}")
    
    # 这里你可以调用新闻API或网页搜索
    # 示例：使用简化的模拟流程
    article_data = {
        "title": f"今日焦点：{topic}最新进展",
        "content": f"这是关于{topic}的自动生成内容...",
        "topic": topic,
        "date": datetime.now().strftime("%Y年%m月%d日")
    }
    
    log(f"文章标题: {article_data['title']}")
    return article_data


def create_word_document(article_data):
    """创建Word文档到桌面"""
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(article_data['title'], 0)
        
        # 添加内容
        doc.add_paragraph(f"发布时间: {article_data['date']}")
        doc.add_paragraph()
        
        # 正文内容（这里你需要实际的内容生成逻辑）
        content = generate_article_content(article_data['topic'])
        doc.add_paragraph(content)
        
        # 保存到桌面
        filename = f"{article_data['title']}.docx"
        filepath = os.path.join(DESKTOP_PATH, filename)
        doc.save(filepath)
        
        log(f"文档已保存: {filepath}")
        return filepath
        
    except ImportError:
        log("错误: 需要安装 python-docx")
        log("请运行: pip install python-docx")
        return None
    except Exception as e:
        log(f"创建文档失败: {e}")
        return None


def generate_article_content(topic):
    """生成文章内容（示例模板）"""
    # 这里你可以接入AI API来生成内容
    # 或者使用预设模板
    templates = {
        "全球变暖": """
全球变暖正在改变我们的世界

近年来，全球变暖问题日益严重，给人类社会和自然环境带来了深远影响。

一、全球变暖的现状

科学家指出，地球平均气温正在持续上升。冰川融化、海平面上升、极端天气频发...

二、主要原因

温室气体排放是主要原因。工业化进程中，大量燃烧化石燃料...

三、应对措施

1. 减少碳排放
2. 发展清洁能源
3. 植树造林
4. 绿色出行

结语：保护地球，人人有责。
        """,
        "人工智能": """
人工智能：改变未来的力量

人工智能技术正在快速发展，深刻改变着我们的生活方式。

一、AI技术的突破

近年来，大语言模型、计算机视觉等领域取得重大突破...

二、应用场景

1. 医疗健康
2. 教育培训
3. 交通出行
4. 金融服务

三、未来展望

AI将继续深入各行各业，成为推动社会进步的重要力量。

结语：拥抱AI时代，共创美好未来。
        """,
    }
    
    return templates.get(topic, f"这是关于{topic}的详细报道内容...")


def schedule_upload_task(docx_path):
    """创建定时上传任务（6:00）"""
    try:
        # 计算明天的6:00
        tomorrow_6am = datetime.now().replace(hour=6, minute=0, second=0, microsecond=0)
        if tomorrow_6am < datetime.now():
            tomorrow_6am += timedelta(days=1)
        
        log(f"设置定时上传任务: {tomorrow_6am}")
        
        # 创建上传脚本路径
        upload_script = Path(CONTENT_PUBLISHER_PATH) / "upload_task.py"
        
        # 生成上传脚本
        upload_code = f'''#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import subprocess
import os
import sys

CONTENT_PUBLISHER_PATH = r"{CONTENT_PUBLISHER_PATH}"
DOCX_PATH = r"{docx_path}"

def log(msg):
    with open(os.path.join(CONTENT_PUBLISHER_PATH, "upload_task.log"), "a", encoding="utf-8") as f:
        from datetime import datetime
        f.write(f"[{{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}}] {{msg}}\\n")

log("开始定时上传任务...")

# 切换到工作目录
os.chdir(CONTENT_PUBLISHER_PATH)

# 上传平台列表
platforms = ["baijiahao", "zhihu", "toutiao"]

for platform in platforms:
    log(f"上传到 {{platform}}...")
    try:
        result = subprocess.run(
            ["python", "publisher.py", "publish", DOCX_PATH, "-p", platform],
            capture_output=True,
            text=True,
            timeout=300
        )
        log(f"{{platform}} 结果: {{result.returncode}}")
        log(f"{{platform}} 输出: {{result.stdout[-500:]}}")  # 最后500字符
    except Exception as e:
        log(f"{{platform}} 错误: {{e}}")

log("定时上传任务完成")
'''
        
        with open(upload_script, "w", encoding="utf-8") as f:
            f.write(upload_code)
        
        # 使用Windows任务计划程序创建任务
        task_name = "ContentPublisher_Upload"
        
        # 删除旧任务
        subprocess.run(f'schtasks /delete /tn {task_name} /f', shell=True, capture_output=True)
        
        # 创建新任务
        cmd = f'''
        schtasks /create /tn {task_name} /tr "{sys.executable} {upload_script}" /sc once /st 06:00 /f
        '''
        result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
        
        if result.returncode == 0:
            log(f"定时任务创建成功: 明天 6:00")
            return True
        else:
            log(f"创建定时任务失败: {result.stderr}")
            return False
            
    except Exception as e:
        log(f"设置定时任务错误: {e}")
        return False


def main():
    """主流程"""
    log("=" * 50)
    log("晨间新闻自动化任务开始")
    log("=" * 50)
    
    # 1. 启动代理
    if not start_proxy():
        log("启动代理失败，终止任务")
        return False
    
    # 2. 搜索新闻并改写文章
    log("开始搜索新闻...")
    try:
        article_data = search_and_rewrite()
    except Exception as e:
        log(f"搜索新闻失败: {e}")
        close_proxy()
        return False
    
    # 3. 创建Word文档
    log("创建Word文档...")
    docx_path = create_word_document(article_data)
    if not docx_path:
        log("创建文档失败")
        close_proxy()
        return False
    
    # 4. 关闭代理
    close_proxy()
    
    # 5. 设置定时上传任务
    log("设置定时上传任务...")
    if schedule_upload_task(docx_path):
        log("任务设置完成！文章将在明天6:00自动上传")
    else:
        log("设置定时任务失败，请手动上传")
    
    log("=" * 50)
    log("晨间任务完成")
    log("=" * 50)
    
    return True


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        log(f"程序异常: {e}")
        import traceback
        log(traceback.format_exc())
        
        # 确保代理被关闭
        try:
            close_proxy()
        except:
            pass
        
        sys.exit(1)
