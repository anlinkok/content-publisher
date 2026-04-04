#!/usr/bin/env python3
"""
Markdown 转 Word 转换器
"""
import sys
import re
from pathlib import Path

def markdown_to_docx(md_file, output_file=None):
    """将 Markdown 文件转换为 Word 文档"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("请先安装 python-docx: pip install python-docx")
        sys.exit(1)
    
    # 读取 Markdown 文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 创建 Word 文档
    doc = Document()
    
    # 处理每一行
    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 跳过空行
        if not line.strip():
            i += 1
            continue
        
        # 处理标题
        if line.startswith('# '):
            # 一级标题
            p = doc.add_heading(line[2:], level=1)
            i += 1
        elif line.startswith('## '):
            # 二级标题
            p = doc.add_heading(line[3:], level=2)
            i += 1
        elif line.startswith('### '):
            # 三级标题
            p = doc.add_heading(line[4:], level=3)
            i += 1
        elif line.startswith('---'):
            # 分隔线
            doc.add_paragraph('─' * 50)
            i += 1
        elif line.startswith('> '):
            # 引用
            quote_text = line[2:]
            # 检查是否有更多引用行
            j = i + 1
            while j < len(lines) and lines[j].startswith('> '):
                quote_text += '\n' + lines[j][2:]
                j += 1
            p = doc.add_paragraph(quote_text)
            p.style = 'Quote'
            i = j
        elif line.startswith('| ') and '|' in line[1:]:
            # 表格开始
            table_lines = [line]
            j = i + 1
            while j < len(lines) and lines[j].startswith('|'):
                table_lines.append(lines[j])
                j += 1
            
            # 解析表格（跳过分隔行）
            rows = []
            for tl in table_lines:
                if '---' in tl or ':-' in tl:
                    continue
                cells = [c.strip() for c in tl.split('|')[1:-1]]
                if cells:
                    rows.append(cells)
            
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                for row_idx, row_cells in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_cells):
                        table.rows[row_idx].cells[col_idx].text = cell_text
            
            i = j
        elif line.startswith('![ '):
            # 图片
            alt_match = re.search(r'!\[(.+?)\]\((.+?)\)', line)
            if alt_match:
                alt_text = alt_match.group(1)
                p = doc.add_paragraph(f'[图片: {alt_text}]')
            i += 1
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            # 粗体段落（标题类）
            text = line.strip().strip('*')
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(12)
            i += 1
        elif re.match(r'^\d+\.\s+', line):
            # 有序列表
            text = re.sub(r'^\d+\.\s+', '', line)
            doc.add_paragraph(text, style='List Number')
            i += 1
        elif line.startswith('- ') or line.startswith('* '):
            # 无序列表
            text = line[2:]
            doc.add_paragraph(text, style='List Bullet')
            i += 1
        else:
            # 普通段落，处理行内格式
            text = line
            # 处理粗体 **text**
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            # 处理斜体 *text*
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            
            p = doc.add_paragraph(text)
            i += 1
    
    # 保存文档
    if not output_file:
        output_file = md_file.replace('.md', '.docx')
    
    doc.save(output_file)
    print(f"✅ 转换完成: {output_file}")
    return output_file

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python md_to_docx.py <markdown文件> [输出docx文件]")
        sys.exit(1)
    
    md_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    markdown_to_docx(md_file, output_file)
