#!/usr/bin/env python3
"""
Word 文档解析器
支持 .docx 格式，提取文本、图片、样式
"""

import os
import io
import re
import hashlib
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image


@dataclass
class WordImage:
    """Word 中的图片"""
    data: bytes
    format: str
    filename: str
    width: Optional[int] = None
    height: Optional[int] = None
    
    @property
    def hash(self) -> str:
        """计算图片哈希，用于去重"""
        return hashlib.md5(self.data).hexdigest()[:8]
    
    def save(self, output_dir: str) -> str:
        """保存图片到指定目录"""
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        filepath = os.path.join(output_dir, f"{self.hash}.{self.format}")
        
        # 如果文件已存在，直接返回
        if os.path.exists(filepath):
            return filepath
        
        with open(filepath, 'wb') as f:
            f.write(self.data)
        
        return filepath


@dataclass
class WordContent:
    """解析后的 Word 内容"""
    title: str
    content: str  # Markdown 格式
    html_content: str
    images: List[WordImage]
    author: Optional[str] = None
    summary: Optional[str] = None
    tags: List[str] = None
    
    def __post_init__(self):
        if self.tags is None:
            self.tags = []


class WordParser:
    """Word 文档解析器"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
        self.images: List[WordImage] = []
        self.image_index = 0
    
    def parse(self) -> WordContent:
        """解析 Word 文档"""
        # 提取标题（通常是第一个大号字体段落或文档标题）
        title = self._extract_title()
        
        # 提取所有图片
        self.images = self._extract_images()
        
        # 转换内容为 Markdown
        markdown_content = self._to_markdown()
        
        # 转换为 HTML
        html_content = self._to_html()
        
        # 提取摘要（前200字）
        summary = self._extract_summary(markdown_content)
        
        # 自动提取标签
        tags = self._extract_tags(markdown_content)
        
        return WordContent(
            title=title,
            content=markdown_content,
            html_content=html_content,
            images=self.images,
            summary=summary,
            tags=tags
        )
    
    def _extract_title(self) -> str:
        """提取文档标题"""
        # 1. 尝试获取文档属性中的标题
        if self.doc.core_properties.title:
            return self.doc.core_properties.title
        
        # 2. 尝试找第一个大号字体的段落
        for para in self.doc.paragraphs[:5]:
            if para.text.strip():
                # 检查字体大小
                for run in para.runs:
                    if run.font.size and run.font.size.pt > 14:
                        return para.text.strip()
        
        # 3. 使用第一个非空段落
        for para in self.doc.paragraphs:
            if para.text.strip():
                return para.text.strip()
        
        # 4. 使用文件名
        return Path(self.file_path).stem
    
    def _extract_images(self) -> List[WordImage]:
        """提取文档中的所有图片"""
        images = []
        
        # 遍历所有关系（图片存储在关系中）
        rels = self.doc.part.rels
        for rel in rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    
                    # 检测图片格式
                    img_format = self._detect_image_format(image_data)
                    
                    # 生成文件名
                    self.image_index += 1
                    filename = f"image_{self.image_index}.{img_format}"
                    
                    # 尝试获取图片尺寸
                    try:
                        img = Image.open(io.BytesIO(image_data))
                        width, height = img.size
                    except:
                        width, height = None, None
                    
                    images.append(WordImage(
                        data=image_data,
                        format=img_format,
                        filename=filename,
                        width=width,
                        height=height
                    ))
                except Exception as e:
                    print(f"提取图片失败: {e}")
        
        return images
    
    def _detect_image_format(self, data: bytes) -> str:
        """检测图片格式"""
        if data.startswith(b'\xff\xd8\xff'):
            return 'jpg'
        elif data.startswith(b'\x89PNG'):
            return 'png'
        elif data.startswith(b'GIF'):
            return 'gif'
        elif data.startswith(b'BM'):
            return 'bmp'
        elif data.startswith(b'RIFF') and data[8:12] == b'WEBP':
            return 'webp'
        else:
            return 'png'  # 默认
    
    def _to_markdown(self) -> str:
        """转换为 Markdown 格式"""
        lines = []
        image_map = {img.filename: img for img in self.images}
        image_counter = 0
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append('')
                continue
            
            # 判断段落样式
            style_name = para.style.name.lower() if para.style else ''
            
            # 标题
            if 'heading' in style_name or '标题' in style_name:
                level = self._get_heading_level(style_name)
                lines.append(f"{'#' * level} {text}")
            
            # 列表
            elif para._p.getchildren() and para._p.getchildren()[0].tag.endswith('numPr'):
                lines.append(f"1. {text}")
            elif style_name.startswith('list'):
                lines.append(f"- {text}")
            
            # 普通段落
            else:
                # 处理粗体和斜体
                formatted_text = self._format_runs(para)
                lines.append(formatted_text)
        
        # 处理表格
        for table in self.doc.tables:
            lines.append('')
            lines.append(self._table_to_markdown(table))
            lines.append('')
        
        # 替换图片占位符
        content = '\n'.join(lines)
        
        # 插入图片引用
        for i, img in enumerate(self.images, 1):
            # 在适当位置插入图片引用
            placeholder = f"[IMAGE_{i}]"
            content = content.replace(placeholder, f"\n\n![图片{i}](./images/{img.hash}.{img.format})\n\n")
        
        return content
    
    def _to_html(self) -> str:
        """转换为 HTML 格式"""
        html_parts = ['<article>']
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            style_name = para.style.name.lower() if para.style else ''
            
            # 标题
            if 'heading' in style_name or '标题' in style_name:
                level = self._get_heading_level(style_name)
                formatted = self._format_runs_html(para)
                html_parts.append(f"<h{level}>{formatted}</h{level}>")
            
            # 列表
            elif para._p.getchildren() and para._p.getchildren()[0].tag.endswith('numPr'):
                formatted = self._format_runs_html(para)
                html_parts.append(f"<ol><li>{formatted}</li></ol>")
            elif style_name.startswith('list'):
                formatted = self._format_runs_html(para)
                html_parts.append(f"<ul><li>{formatted}</li></ul>")
            
            # 普通段落
            else:
                formatted = self._format_runs_html(para)
                if formatted:
                    html_parts.append(f"<p>{formatted}</p>")
        
        # 表格
        for table in self.doc.tables:
            html_parts.append(self._table_to_html(table))
        
        html_parts.append('</article>')
        return '\n'.join(html_parts)
    
    def _get_heading_level(self, style_name: str) -> int:
        """获取标题级别"""
        if 'heading 1' in style_name or '标题 1' in style_name:
            return 1
        elif 'heading 2' in style_name or '标题 2' in style_name:
            return 2
        elif 'heading 3' in style_name or '标题 3' in style_name:
            return 3
        elif 'heading 4' in style_name or '标题 4' in style_name:
            return 4
        elif 'heading 5' in style_name or '标题 5' in style_name:
            return 5
        elif 'heading 6' in style_name or '标题 6' in style_name:
            return 6
        return 2  # 默认二级标题
    
    def _format_runs(self, para) -> str:
        """格式化段落中的 runs（粗体、斜体等）"""
        parts = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            # 处理样式
            if run.bold:
                text = f"**{text}**"
            if run.italic:
                text = f"*{text}*"
            if run.underline:
                text = f"<u>{text}</u>"
            
            parts.append(text)
        
        return ''.join(parts)
    
    def _format_runs_html(self, para) -> str:
        """格式化段落中的 runs 为 HTML"""
        parts = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            # 转义 HTML 特殊字符
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # 处理样式
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            
            parts.append(text)
        
        return ''.join(parts)
    
    def _table_to_markdown(self, table) -> str:
        """表格转 Markdown"""
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows.append('| ' + ' | '.join(cells) + ' |')
            
            # 添加分隔符
            if i == 0:
                rows.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
        
        return '\n'.join(rows)
    
    def _table_to_html(self, table) -> str:
        """表格转 HTML"""
        html = ['<table border="1" cellpadding="5">']
        
        for i, row in enumerate(table.rows):
            html.append('<tr>')
            tag = 'th' if i == 0 else 'td'
            for cell in row.cells:
                html.append(f'<{tag}>{cell.text.strip()}</{tag}>')
            html.append('</tr>')
        
        html.append('</table>')
        return '\n'.join(html)
    
    def _extract_summary(self, content: str, max_length: int = 200) -> str:
        """提取摘要"""
        # 移除 Markdown 标记
        text = re.sub(r'[#*`\[\]!]', '', content)
        text = re.sub(r'\n+', ' ', text)
        text = text.strip()
        
        if len(text) <= max_length:
            return text
        
        return text[:max_length] + '...'
    
    def _extract_tags(self, content: str) -> List[str]:
        """提取标签"""
        # 简单的关键词提取
        keywords = {
            '科技': ['AI', '人工智能', '技术', '编程', '软件', '硬件'],
            '互联网': ['互联网', '产品', '运营', '数据', '用户'],
            '财经': ['股票', '投资', '理财', '经济', '金融'],
            '生活': ['生活', '美食', '旅行', '健康'],
            '职场': ['职场', '管理', '沟通', '效率'],
            '教育': ['学习', '教育', '知识', '读书'],
        }
        
        tags = []
        for category, words in keywords.items():
            for word in words:
                if word in content:
                    tags.append(category)
                    break
        
        return tags[:5]
    
    def save_images(self, output_dir: str) -> List[str]:
        """保存所有图片到指定目录"""
        paths = []
        for img in self.images:
            path = img.save(output_dir)
            paths.append(path)
        return paths


# ============================================================================
# CLI 工具
# ============================================================================

if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='Word 文档解析器')
    parser.add_argument('file', help='Word 文件路径')
    parser.add_argument('--output', '-o', help='输出 Markdown 文件路径')
    parser.add_argument('--images', '-i', help='图片输出目录')
    
    args = parser.parse_args()
    
    if not os.path.exists(args.file):
        print(f"文件不存在: {args.file}")
        exit(1)
    
    # 解析文档
    parser = WordParser(args.file)
    content = parser.parse()
    
    print(f"标题: {content.title}")
    print(f"图片数: {len(content.images)}")
    print(f"标签: {', '.join(content.tags)}")
    print(f"\n内容预览:\n{content.content[:500]}...")
    
    # 保存 Markdown
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(f"---\n")
            f.write(f"title: \"{content.title}\"\n")
            f.write(f"tags: {content.tags}\n")
            f.write(f"---\n\n")
            f.write(content.content)
        print(f"\n已保存到: {args.output}")
    
    # 保存图片
    if args.images:
        paths = parser.save_images(args.images)
        print(f"已保存 {len(paths)} 张图片到: {args.images}")
